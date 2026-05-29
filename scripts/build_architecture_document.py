#!/usr/bin/env python3

from __future__ import annotations

import json
import sqlite3
from dataclasses import dataclass
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "data" / "app.db"
SCREENSHOT_DIR = ROOT / "artifacts" / "architecture" / "screenshots"
DOC_OUTPUT = ROOT / "output" / "doc" / "RRCP_Architecture_Working_Guide.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "RRCP_Architecture_Working_Guide.pdf"


SCREENSHOTS = [
    ("Overview", "overview.png", "Monthly workspace summary, current status, and navigation starting point."),
    ("Setup Center", "setup-center.png", "Master uploads, doctor ownership, user access, and approval request forms."),
    ("Monthly Intake", "monthly-intake.png", "Monthly transaction upload, parsed rows, payment method, and revenue booked in."),
    ("Calculation Review", "calculation-review.png", "Engine output with allowed discount, actual discount, incentive, variance, and remarks."),
    ("Payout Center", "payout-center.png", "Payment entries, approvals, cash-in-hand control, and period lock workflow."),
    ("Reports", "reports.png", "Exportable reports for doctor, PRO, group, and operational review."),
    ("Support", "support.png", "Support requests and operational issue tracking."),
]

SYSTEM_TREE = """RRCP Referral Revenue Calculation Platform
├─ Login
│  ├─ Input: email and password
│  ├─ API: POST /api/auth/login
│  ├─ Reads: users, doctor_master
│  └─ Output: secure token, role, linked doctor profile
├─ Overview
│  ├─ API: GET /api/dashboard
│  ├─ Reads: transactions, engine_runs, payments, approval_requests, reference tables
│  └─ Output: monthly KPIs, pending work, latest run status, shortcut cards
├─ Setup Center
│  ├─ Upload Master Data
│  │  ├─ API: POST /api/reference/upload
│  │  ├─ Reads file sheets: service/item sheet, doctor group sheet
│  │  ├─ Writes: service_prices, discount_rules, doctor_master, reference_uploads
│  │  └─ Output: doctor list, item list, discount rules, group rules
│  ├─ Upload Rule Notes
│  │  ├─ API: POST /api/reference/software/upload
│  │  ├─ Writes: software_requirements, reference_uploads
│  │  └─ Output: business-rule checklist used by users during review
│  ├─ Doctor Ownership and Mapping
│  │  ├─ API: GET /api/reference/doctors
│  │  ├─ Reads: doctor_master
│  │  └─ Output: doctor profile, group, PRO, cycle, reporting doctor, verification state
│  ├─ Update Doctor Info
│  │  ├─ API: POST /api/reference/doctor-info-change
│  │  ├─ Writes: approval_requests
│  │  └─ Output: pending approval before master changes
│  ├─ Add Doctor
│  │  ├─ API: POST /api/reference/doctor-addition
│  │  ├─ Writes: approval_requests first, doctor_master after approval
│  │  └─ Output: controlled doctor creation
│  └─ User Management
│     ├─ API: GET/POST /api/users
│     ├─ Writes: users
│     └─ Output: admin, mapper, accountant, doctor access
├─ Monthly Intake
│  ├─ Upload Monthly File
│  │  ├─ API: POST /api/data/upload
│  │  ├─ Reads file columns: patient, doctor, PRO, item, price, discount, payment, booked-in, incentive
│  │  ├─ Writes: transactions, transaction_incentives, reference_uploads
│  │  └─ Output: parsed monthly rows for selected period
│  ├─ Search/Filter Table
│  │  ├─ API: GET /api/data
│  │  ├─ Reads: transactions
│  │  └─ Output: imported rows with patient, payment, method, booked-in, and status columns
│  └─ Download Current Period CSV
│     ├─ API: GET /api/data/export
│     └─ Output: CSV of filtered imported data
├─ Calculation Review
│  ├─ Run Engine
│  │  ├─ API: POST /api/engine/run
│  │  ├─ Reads: transactions, doctor_master, discount_rules, transaction_incentives
│  │  ├─ Writes: engine_runs, engine_results
│  │  └─ Output: allowed amount, actual discount, incentive, variance, remark
│  ├─ Search/Filter Result Table
│  │  ├─ API: GET /api/engine/results
│  │  └─ Output: row-level calculation explanation
│  └─ Request Incentive Override
│     ├─ API: POST /api/engine/override-request
│     ├─ Writes: approval_requests
│     └─ Output: pending approval for changed payable amount
├─ Payout Center
│  ├─ Generate/Review Payments
│  │  ├─ API: GET/POST /api/payments
│  │  ├─ Reads: engine_results, doctor_master, approval_requests
│  │  ├─ Writes: payments, approval_requests
│  │  └─ Output: payable list and disbursal approvals
│  ├─ Cash In Hand Control
│  │  ├─ Reads/Writes: payments and wallet/cash tracking fields
│  │  └─ Output: blocks fresh disbursal until hand amount is zero
│  └─ Lock Period
│     ├─ Writes: locked_periods
│     └─ Output: closed month so data cannot be changed accidentally
├─ Reports
│  ├─ API: GET /api/reports/*
│  ├─ Reads: transactions, engine_results, payments, doctor_master
│  └─ Output: doctor, PRO, group, exception, and payout reports
└─ Support
   ├─ API: POST /api/contact and GET /api/contact/messages
   ├─ Writes: contact_messages
   └─ Output: user support/issues list"""

PAGE_ACTIONS = [
    ["Login", "Login button", "POST /api/auth/login", "users, doctor_master", "Validates credentials and opens the app with the correct role and linked doctor access."],
    ["Header", "Logout button", "Browser local token/session", "No database write", "Clears current login and returns user to login screen."],
    ["Left navigation", "Overview", "GET /api/dashboard", "transactions, engine_runs, payments, approvals", "Shows monthly summary and pending operational work."],
    ["Left navigation", "Setup Center", "GET /api/reference/summary, GET /api/users, GET /api/reference/doctors", "reference tables, users", "Opens master upload, doctor mapping, doctor requests, and user creation."],
    ["Left navigation", "Monthly Intake", "GET /api/data", "transactions", "Opens monthly upload and imported transaction table."],
    ["Left navigation", "Calculation Review", "GET /api/engine/results", "engine_runs, engine_results", "Shows latest calculation output and exceptions."],
    ["Left navigation", "Payout Center", "GET /api/payments", "payments, engine_results, approvals", "Shows payable/disbursal workflow."],
    ["Left navigation", "Reports", "GET /api/reports/*", "transactions, engine_results, payments", "Shows exportable business reports."],
    ["Left navigation", "Support", "GET /api/contact/messages", "contact_messages", "Shows support requests and contact form."],
    ["Setup Center", "Download Sample Format", "Static/template file", "No database write", "Downloads the expected Excel format for correct upload columns."],
    ["Setup Center", "Update Master Data", "POST /api/reference/upload", "service_prices, discount_rules, doctor_master, reference_uploads", "Imports doctor master, doctor groups, item rules, and allowed discounts."],
    ["Setup Center", "Update Rule Notes", "POST /api/reference/software/upload", "software_requirements, reference_uploads", "Imports business requirement notes so users know what rules to follow."],
    ["Setup Center", "Search doctor", "GET /api/reference/doctors?search=", "doctor_master", "Finds a doctor by name, code, PRO, group, degree, contact number, cycle, or confirmation status."],
    ["Setup Center", "Use In Request", "UI state plus doctor_master row", "No database write until request submitted", "Copies selected doctor details into the change request form."],
    ["Setup Center", "Verify/Unverify", "Doctor verification endpoint/update flow", "doctor_master", "Marks whether the doctor has been operationally verified."],
    ["Setup Center", "Submit doctor info change", "POST /api/reference/doctor-info-change", "approval_requests", "Creates approval before changing group, PRO, cycle, reporting doctor, confirmation status, degree, or contact."],
    ["Setup Center", "Submit new doctor", "POST /api/reference/doctor-addition", "approval_requests, later doctor_master", "Creates a doctor addition request; doctor is inserted only after approval."],
    ["Setup Center", "Create User", "POST /api/users", "users", "Creates login access and optionally links a doctor user to a doctor profile."],
    ["Monthly Intake", "Choose file", "Browser file picker", "No database write", "Selects the monthly transaction or incentive workbook from the computer."],
    ["Monthly Intake", "Upload", "POST /api/data/upload", "transactions, transaction_incentives, reference_uploads", "Parses the workbook, removes duplicate rows using row-level matching, and stores monthly data."],
    ["Monthly Intake", "Search table", "GET /api/data?search=", "transactions", "Filters imported rows across patient, doctor, item, PRO, payment method, booked-in, amounts, status, and dates."],
    ["Monthly Intake", "Download Current Period CSV", "GET /api/data/export", "transactions", "Exports the visible monthly transaction data for external checking."],
    ["Calculation Review", "Run Engine", "POST /api/engine/run", "transactions, doctor_master, discount_rules, transaction_incentives, engine_runs, engine_results", "Creates a new run and calculates every row for the selected period."],
    ["Calculation Review", "Search table", "GET /api/engine/results?search=", "engine_results", "Filters calculation rows by doctor, group, PRO, item, remark, amount, and status."],
    ["Calculation Review", "Request Override", "POST /api/engine/override-request", "approval_requests", "Creates approval if incentive/payable amount needs manual override."],
    ["Payout Center", "Refresh/Load payments", "GET /api/payments", "payments, engine_results", "Loads payout rows for the current period."],
    ["Payout Center", "Create/Submit disbursal", "POST /api/payments or approval request", "payments, approval_requests", "Moves eligible calculated incentives into payment workflow."],
    ["Payout Center", "Approve/Reject disbursal", "PATCH /api/approvals/{id}", "approval_requests, payments", "Approves or rejects final payment release."],
    ["Payout Center", "Lock Period", "Period lock endpoint", "locked_periods", "Prevents accidental changes after monthly close."],
    ["Approvals", "Approve", "PATCH /api/approvals/{id}", "approval_requests plus target table", "Applies the approved change to doctor_master, payments, or payout status."],
    ["Approvals", "Reject", "PATCH /api/approvals/{id}", "approval_requests", "Keeps original data unchanged and records rejected state."],
    ["Reports", "Download/Export", "GET /api/reports/*", "transactions, engine_results, payments", "Exports doctor, PRO, group, exception, or payout views."],
    ["Support", "Send message", "POST /api/contact", "contact_messages", "Stores a support request for admin review."],
]

DATA_LINEAGE = [
    ["Users and roles", "Setup Center create user or seeded admin/demo users", "users", "Login, role-based access, doctor-only data filtering"],
    ["Linked doctor for doctor login", "Create User linked doctor profile", "users.doctor_master_id -> doctor_master.id", "Doctor users can see only their own doctor-related data"],
    ["Doctor name", "Special Discount Master doctor sheet or Add Doctor approval", "doctor_master.doctor_name, transactions.referring_doctor", "Used to match transaction row to doctor master"],
    ["Doctor code", "Doctor master upload/add doctor request", "doctor_master.doctor_code", "Used for identification and dropdown labels"],
    ["Degree and contact no.", "Doctor master upload or doctor info change approval", "doctor_master.degree, doctor_master.contact_no", "Shown in doctor profile and search"],
    ["Doctor group", "Doctor master upload or approved doctor info change", "doctor_master.incentive_group", "Drives Group A/B/C/D/E/F rule selection"],
    ["PRO ownership", "Doctor master upload or approved PRO change", "doctor_master.present_pro, transactions.pro_name", "Used for ownership, reports, and payout routing"],
    ["Incentive cycle", "Doctor master upload or doctor info change", "doctor_master.incentive_cycle", "Shows payment cycle timing"],
    ["Reporting doctor", "Doctor master upload or doctor info change", "doctor_master.reporting_doctor", "Used for hierarchy/reporting review"],
    ["Confirmation status", "Doctor master upload or doctor info change", "doctor_master.confirmation_status, confirmation_remarks", "Shows confirmed/pending/no confirmation state"],
    ["Verified doctor", "Doctor mapping verify action", "doctor_master.verified", "Shows whether the doctor has been checked before monthly processing"],
    ["Item/modality", "Monthly transaction upload and Special Discount Master", "transactions.billable_items, transactions.modality, discount_rules.normalized_item", "Required to find allowed discount/incentive rule"],
    ["Patient ID and patient name", "Monthly transaction upload", "transactions.patient_id, transactions.patient_name", "Used for row traceability and duplicate detection"],
    ["Total payment received", "Monthly transaction upload", "transactions.total_payment", "Displayed in Monthly Intake and exports"],
    ["Payment method", "Monthly transaction upload", "transactions.payment_method", "Displayed in intake table and searchable"],
    ["Revenue booked in", "Monthly transaction upload", "transactions.revenue_booked_in", "Shows Sukhmani/Jivada booking source"],
    ["Total discount amount", "Monthly transaction upload", "transactions.total_discount", "Actual discount used in variance calculation"],
    ["Incentive to doctors", "Monthly incentive workbook column or transaction incentive columns", "transaction_incentives.incentive_amount/payable_discount or fallback", "Used in Sum of Both and payout calculation"],
    ["Allowed discount", "Special Discount Master item/group rules", "discount_rules.group_json", "Target allowed value for each doctor group and item"],
    ["Calculation run", "Run Engine button", "engine_runs", "Stores run id, period, totals, and summary JSON"],
    ["Calculation row output", "Run Engine button", "engine_results", "Stores allowed, actual, incentive, variance, remark, approval required"],
    ["Approval request", "Doctor change, PRO change, override, disbursal, add doctor", "approval_requests", "Controls changes that should not happen directly"],
    ["Payment/disbursal", "Payout Center workflow", "payments", "Tracks payout amount, status, approval, adjustment, advance, and cash handover state"],
    ["Support request", "Support page contact form", "contact_messages", "Stores operational issues raised by users"],
]

CALCULATION_TRACE = [
    ["1", "Select period", "The current month/year selector defines the date range.", "transactions.visit_date"],
    ["2", "Load monthly rows", "Rows uploaded in Monthly Intake are pulled for that date range.", "transactions"],
    ["3", "Normalize doctor and item", "Names are cleaned to reduce mismatch from spaces, punctuation, and case differences.", "normalized_doctor, normalized_item"],
    ["4", "Find doctor master", "The engine searches doctor_master by normalized doctor name.", "doctor_master"],
    ["5", "Find group and PRO", "If doctor exists, group, PRO, cycle, reporting doctor, confirmation and verification state are read.", "doctor_master"],
    ["6", "Find item rule", "The engine searches discount_rules by billable item or modality.", "discount_rules"],
    ["7", "Read exact incentive", "If a matching incentive workbook row exists, that imported value is used before fallback.", "transaction_incentives"],
    ["8", "Apply group policy", "Allowed value is selected by group. Group A must be zero. Group B uses Group D discount and zero incentive. Group C uses Group F discount and zero incentive.", "discount_rules.group_json and doctor_master.incentive_group"],
    ["9", "Calculate sum", "Sum Of Both equals Total Discount Amount plus Incentive To Doctors.", "engine_results.actual_discount + incentive_amount"],
    ["10", "Calculate variance", "Variance equals Allowed minus Sum Of Both.", "engine_results.variance_amount"],
    ["11", "Create remark", "Missing doctor, missing group, missing item, over-discount, lower discount, OK, or group mismatch text is stored.", "engine_results.remark"],
    ["12", "Mark approval", "Rows needing manual approval are marked so Payout Center/Approvals can control the next step.", "engine_results.approval_required, approval_requests"],
]

REMARK_EXPLAINER = [
    ["OK", "Doctor, group, item rule, actual discount, and incentive are matching the expected rule.", "Can proceed to payout review."],
    ["Doctor name missing in master sheet", "The transaction doctor could not be matched with doctor_master.", "Add/approve doctor or correct doctor name in master/upload."],
    ["Doctor group missing in master sheet", "Doctor exists but group is blank or unavailable.", "Submit doctor info change and approve group update."],
    ["Need to master sheet for item", "The billable item/modality is not present in discount_rules.", "Update Special Discount Master with item/group rule and re-run engine."],
    ["Over-discount requires approval", "Total Discount Amount plus Incentive To Doctors is higher than allowed.", "Review row and approve override only if business approves."],
    ["Lower discount than allowed", "Actual discount/incentive is lower than allowed.", "Usually informational, review if payout should be adjusted."],
    ["Group A rule mismatch", "Group A should have zero discount and zero incentive.", "Fix transaction/incentive or approve exception."],
    ["Group B rule mismatch", "Group B should use Group D discount and zero doctor incentive.", "Fix incentive/discount or approve exception."],
    ["Group C rule mismatch", "Group C should use Group F discount and zero doctor incentive.", "Fix incentive/discount or approve exception."],
]

ROLE_GUIDE = [
    ["Admin", "Full setup, uploads, run engine, approvals, payout control, reports, user creation.", "Owns master accuracy and final monthly close."],
    ["Mapper", "Doctor mapping, doctor info requests, PRO/group/cycle verification.", "Prepares clean master data before calculation."],
    ["Accountant", "Payout review, disbursal checks, payment adjustments, cash-in-hand tracking, reports.", "Owns payment correctness after calculation."],
    ["Doctor", "Doctor-specific view only when linked to doctor profile.", "Can see own data needed for confirmation/review."],
]

MONTHLY_WORKFLOW_STEPS = [
    {
        "step": "Step 1",
        "title": "Prepare master data before starting the month",
        "screen": "Setup Center",
        "starts_from": "The process starts with reference files, not with calculation. Admin uploads the latest Special Discount Master and software requirement/rule notes.",
        "input": "Special Discount Master and Software Requirement file.",
        "data_captured": "Doctor names, doctor codes, degrees, contact numbers, PRO ownership, incentive groups, incentive cycles, reporting doctors, confirmation state, billable items, modalities, allowed discount values, and business rule notes.",
        "stored_in": "doctor_master, service_prices, discount_rules, software_requirements, reference_uploads.",
        "why": "The engine cannot calculate correctly unless it knows which doctor belongs to which group and which item has which allowed discount rule.",
        "output": "Clean master tables are ready for doctor verification and monthly transaction matching.",
        "screenshot": "setup-center.png",
    },
    {
        "step": "Step 2",
        "title": "Verify doctor ownership and doctor profile",
        "screen": "Setup Center",
        "starts_from": "After master upload, the doctor list is searched and reviewed. The user checks whether each doctor has the correct group, PRO, incentive cycle, reporting doctor, degree, contact number, confirmation status, and verified status.",
        "input": "Doctor records created from master upload or approved doctor addition requests.",
        "data_captured": "Doctor profile status, confirmation remarks, verification status, group, PRO, reporting doctor, and contact information.",
        "stored_in": "doctor_master directly for verified fields where allowed, or approval_requests first when the change requires approval.",
        "why": "Wrong doctor group or wrong PRO will produce wrong allowed amount, wrong incentive, wrong payout owner, and wrong report totals.",
        "output": "Doctors are ready to be matched with monthly transactions.",
        "screenshot": "setup-center.png",
    },
    {
        "step": "Step 3",
        "title": "Create user access for operational roles",
        "screen": "Setup Center",
        "starts_from": "Admin creates logins for admin, mapper, accountant, and doctor users. Doctor users are linked to one doctor profile.",
        "input": "Email, temporary password, role, and linked doctor profile if role is Doctor.",
        "data_captured": "User email, password hash, role, active status, linked doctor id.",
        "stored_in": "users.",
        "why": "The system needs role-based control so doctor users can see their own data while admin/accountant can manage calculation, approvals, and payouts.",
        "output": "Users can log in and work only within their permitted responsibility.",
        "screenshot": "setup-center.png",
    },
    {
        "step": "Step 4",
        "title": "Upload monthly transaction and incentive data",
        "screen": "Monthly Intake",
        "starts_from": "Once master data is ready, the monthly transaction file is uploaded for the selected period. If the workbook contains incentive columns, those incentive rows are captured at the same time.",
        "input": "Monthly Dashboard/Incentive workbook for the selected month.",
        "data_captured": "Patient id, patient name, visit date, referring doctor, PRO, billable item, modality, price, total discount amount, total net, total payment received, payment method, revenue booked in Sukhmani/Jivada, receipt status, and imported doctor incentive fields.",
        "stored_in": "transactions, transaction_incentives, reference_uploads.",
        "why": "This is the actual month workload. These rows are what the engine later checks against master rules.",
        "output": "A searchable monthly intake table appears with all imported patient/doctor/payment rows.",
        "screenshot": "monthly-intake.png",
    },
    {
        "step": "Step 5",
        "title": "Check imported data before calculation",
        "screen": "Monthly Intake",
        "starts_from": "Before running the engine, the user searches and reviews imported rows. The purpose is to catch missing doctor names, missing payment values, wrong PRO text, wrong booked-in source, or missing item names before calculation.",
        "input": "Rows already stored from the monthly upload.",
        "data_captured": "No new calculation data is created in this review step unless another corrected file is uploaded.",
        "stored_in": "transactions remains the source of truth for imported monthly rows.",
        "why": "The cleaner the intake data is, the fewer exceptions appear in Calculation Review.",
        "output": "User confirms monthly data is ready to run or uploads a corrected file.",
        "screenshot": "monthly-intake.png",
    },
    {
        "step": "Step 6",
        "title": "Run the calculation engine",
        "screen": "Calculation Review",
        "starts_from": "Admin clicks Run Engine for the selected period. The engine reads transaction rows, doctor master, discount rules, and imported incentive rows.",
        "input": "transactions plus doctor_master plus discount_rules plus transaction_incentives.",
        "data_captured": "Run id, run period, result rows, allowed amount, total discount amount, incentive to doctors, sum of both, variance, remarks, and approval-required state.",
        "stored_in": "engine_runs and engine_results.",
        "why": "This is where the system converts raw monthly rows into rule-checked payout calculation rows.",
        "output": "A calculation table is created showing OK rows, missing master rows, over-discount rows, group policy mismatches, and rows needing approval.",
        "screenshot": "calculation-review.png",
    },
    {
        "step": "Step 7",
        "title": "Understand why each result row is shown",
        "screen": "Calculation Review",
        "starts_from": "The user reviews the engine result table row by row. Each row explains what the system found and why it marked the row OK or flagged it.",
        "input": "engine_results for the latest run.",
        "data_captured": "No new row is created during viewing. If an override is needed, an approval request is created.",
        "stored_in": "engine_results for calculated rows and approval_requests for override requests.",
        "why": "This step explains whether the row can move to payout or whether master data/approval is required first.",
        "output": "Rows are either accepted, corrected through master updates, or sent for approval.",
        "screenshot": "calculation-review.png",
    },
    {
        "step": "Step 8",
        "title": "Resolve exceptions and approval-required items",
        "screen": "Approvals and Setup Center",
        "starts_from": "If a calculation row shows missing doctor, missing group, wrong PRO, group mismatch, over-discount, incentive override, or disbursal request, it must be corrected or approved.",
        "input": "Pending approval requests and flagged engine result rows.",
        "data_captured": "Requested change payload, request type, requested by, approved/rejected status, approved by, updated timestamp.",
        "stored_in": "approval_requests and then target tables like doctor_master or payments after approval.",
        "why": "The system prevents silent changes to doctor information, PRO ownership, incentive amount, doctor addition, and disbursal.",
        "output": "Approved requests update the correct target data. Rejected requests leave original data unchanged.",
        "screenshot": "setup-center.png",
    },
    {
        "step": "Step 9",
        "title": "Prepare payout and disbursal",
        "screen": "Payout Center",
        "starts_from": "After calculation rows are clean or approved, accountant/admin reviews payout entries. Cash-in-hand, advance, adjustment, return of incentive, and handover delay controls are checked here.",
        "input": "Approved engine result rows and payout/payment records.",
        "data_captured": "Payable amount, adjustment, advance payment, approval status, payment status, disbursal notes, handover/cash-in-hand tracking.",
        "stored_in": "payments, approval_requests, locked_periods where period close is used.",
        "why": "No fresh disbursal should happen if cash-in-hand is not zero or if approval is pending.",
        "output": "Final payout/disbursal list is ready for payment or held until issues are cleared.",
        "screenshot": "payout-center.png",
    },
    {
        "step": "Step 10",
        "title": "Export reports and close the month",
        "screen": "Reports",
        "starts_from": "After payout review, reports are exported for management, doctor/PRO review, and accounting records. The period can then be locked.",
        "input": "transactions, engine_results, payments, doctor_master, approval history.",
        "data_captured": "No new business calculation is created by exports. Locking creates a period lock record.",
        "stored_in": "Reports read existing tables. Month close uses locked_periods.",
        "why": "Reports provide the final trace from imported transaction to calculated output and payment decision.",
        "output": "Final reports, exception summaries, payout records, and a closed period.",
        "screenshot": "reports.png",
    },
]


@dataclass
class DbSnapshot:
    counts: list[tuple[str, int]]
    latest_uploads: list[tuple[str, str, int, str]]
    latest_run: dict
    latest_remarks: list[tuple[str, int]]
    approvals: list[tuple[str, str, int]]


def fetch_snapshot() -> DbSnapshot:
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        counts = []
        for table in [
            "users",
            "doctor_master",
            "service_prices",
            "discount_rules",
            "software_requirements",
            "transactions",
            "transaction_incentives",
            "engine_runs",
            "engine_results",
            "payments",
            "approval_requests",
            "locked_periods",
            "contact_messages",
        ]:
            count = conn.execute(f"SELECT COUNT(*) AS count FROM {table}").fetchone()["count"]
            counts.append((table, int(count)))

        latest_uploads = [
            (row["type"], row["file_name"], int(row["row_count"]), row["uploaded_at"])
            for row in conn.execute(
                "SELECT type, file_name, row_count, uploaded_at FROM reference_uploads ORDER BY id DESC LIMIT 8"
            ).fetchall()
        ]

        latest_run_row = conn.execute(
            "SELECT * FROM engine_runs ORDER BY id DESC LIMIT 1"
        ).fetchone()
        latest_run = dict(latest_run_row) if latest_run_row else {}
        if latest_run.get("summary_json"):
            try:
                latest_run["summary"] = json.loads(latest_run["summary_json"])
            except json.JSONDecodeError:
                latest_run["summary"] = {}
        else:
            latest_run["summary"] = {}

        run_id = int(latest_run.get("id") or 0)
        latest_remarks = []
        if run_id:
            latest_remarks = [
                (row["remark"] or "-", int(row["rows"]))
                for row in conn.execute(
                    "SELECT remark, COUNT(*) AS rows FROM engine_results WHERE run_id = ? GROUP BY remark ORDER BY rows DESC",
                    (run_id,),
                ).fetchall()
            ]

        approvals = [
            (row["type"], row["status"], int(row["rows"]))
            for row in conn.execute(
                "SELECT type, status, COUNT(*) AS rows FROM approval_requests GROUP BY type, status ORDER BY type, status"
            ).fetchall()
        ]

    return DbSnapshot(
        counts=counts,
        latest_uploads=latest_uploads,
        latest_run=latest_run,
        latest_remarks=latest_remarks,
        approvals=approvals,
    )


def add_doc_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_doc_bullets(document: Document, bullets: list[str]) -> None:
    for item in bullets:
        document.add_paragraph(item, style="List Bullet")


def add_doc_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def add_doc_preformatted(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)


def add_doc_image(document: Document, image_path: Path, width_inches: float = 9.2) -> None:
    if image_path.exists():
        document.add_picture(str(image_path), width=Inches(width_inches))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_doc_workflow_steps(document: Document) -> None:
    for item in MONTHLY_WORKFLOW_STEPS:
        add_doc_heading(document, f"{item['step']}: {item['title']}", 2)
        add_doc_table(
            document,
            ["Part", "Explanation"],
            [
                ["Screen used", item["screen"]],
                ["Where this step starts", item["starts_from"]],
                ["Input used", item["input"]],
                ["Data captured", item["data_captured"]],
                ["Stored in", item["stored_in"]],
                ["Why this step is needed", item["why"]],
                ["Output of this step", item["output"]],
            ],
        )
        add_doc_image(document, SCREENSHOT_DIR / item["screenshot"], width_inches=8.7)


def build_docx(snapshot: DbSnapshot) -> None:
    DOC_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor(25, 43, 84)

    title = document.add_heading("RRCP Architecture and Working Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = document.add_paragraph(
        "End-to-end explanation of the Referral Revenue Calculation Platform, including imports, captured data, calculation logic, approvals, outputs, and screenshots from the current application."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_doc_heading(document, "1. Tool Purpose", 1)
    add_doc_bullets(
        document,
        [
            "The tool manages monthly doctor referral discount and incentive calculation.",
            "It imports transaction files, reference masters, and incentive mappings.",
            "It validates doctor, group, PRO, and item rules before producing payout-ready results.",
            "It supports approval workflows for doctor changes, PRO changes, doctor additions, incentive overrides, and disbursal.",
            "It exports reports and tracks period locks for month-end closure.",
        ],
    )

    add_doc_heading(document, "2. Code Stack", 1)
    add_doc_table(
        document,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "HTML, CSS, vanilla JavaScript", "Single-page operational UI and table workflows"],
            ["Backend", "PHP", "API routes, upload handling, calculation workflow, approvals"],
            ["Database", "SQLite", "Local persistent storage for users, masters, transactions, runs, payouts"],
            ["Spreadsheet import", "Composer PHP spreadsheet library", "Reads XLSX/XLS files for masters and transaction uploads"],
            ["Artifacts", "Python scripts", "Screenshots, PPT/PDF/DOCX documentation, utility scripts"],
            ["Local runtime", "XAMPP", "Local PHP/Apache hosting at http://localhost/doctor"],
        ],
    )

    add_doc_heading(document, "3. Current Live Data Snapshot", 1)
    add_doc_table(document, ["Table", "Rows"], [[name, str(count)] for name, count in snapshot.counts])

    latest_run = snapshot.latest_run
    summary = latest_run.get("summary", {})
    add_doc_heading(document, "4. Current Engine Run Used In This Guide", 1)
    add_doc_table(
        document,
        ["Field", "Value"],
        [
            ["Run ID", str(latest_run.get("id", "-"))],
            ["Period", f"{latest_run.get('period_year', '-')}-{str(latest_run.get('period_month', '-')).zfill(2)}"],
            ["Total records", str(summary.get("totalRecords", latest_run.get("total_records", "-")))],
            ["Total flags", str(summary.get("totalFlags", latest_run.get("total_flags", "-")))],
            ["Over-discount count", str(summary.get("overDiscountCount", "-"))],
            ["Missing doctor count", str(summary.get("missingDoctorCount", "-"))],
            ["Missing group count", str(summary.get("missingGroupCount", "-"))],
            ["Missing item count", str(summary.get("missingItemCount", "-"))],
            ["Total payable incentive", f"INR {summary.get('totalPayable', '-')}"],
            ["Run timestamp", str(latest_run.get("run_at", "-"))],
        ],
    )

    add_doc_heading(document, "5. Data Import Sources", 1)
    add_doc_table(
        document,
        ["Upload Type", "File", "Rows", "Uploaded At"],
        [[typ, name, str(rows), at] for typ, name, rows, at in snapshot.latest_uploads],
    )

    add_doc_heading(document, "6. Data Flow", 1)
    add_doc_bullets(
        document,
        [
            "Setup Center imports Special Discount Master into service_prices, discount_rules, and doctor_master.",
            "Setup Center imports Software Requirement notes into software_requirements.",
            "Monthly Intake imports transaction rows into transactions and exact incentive mapping rows into transaction_incentives when present.",
            "Calculation Review reads transactions, doctor_master, discount_rules, and transaction_incentives to create an engine_runs record and engine_results rows.",
            "Payout Center creates payments from approved engine results and applies approval/cash-in-hand constraints.",
            "Reports read engine_results, payments, doctors, and transactions for exports.",
        ],
    )

    add_doc_heading(document, "7. Step-by-Step Working From Start to Final Output", 1)
    document.add_paragraph(
        "Read this section first. It explains the tool in the same order a monthly user should operate it: prepare master data, verify doctors, upload monthly data, run calculation, resolve exceptions, approve, payout, report, and close."
    )
    add_doc_workflow_steps(document)

    add_doc_heading(document, "8. Tree View of the Same Flow", 1)
    document.add_paragraph(
        "This is the same process in tree format. It is included after the workflow so the reader can see how the pages, APIs, tables, and outputs connect."
    )
    add_doc_preformatted(document, SYSTEM_TREE)

    add_doc_heading(document, "9. Field Data Lineage", 1)
    document.add_paragraph(
        "This section answers: where does this field come from, where is it stored, and why does the tool show it."
    )
    add_doc_table(document, ["Data Shown", "Fetched or Captured From", "Stored In", "Why It Is Needed"], DATA_LINEAGE)

    add_doc_heading(document, "10. Captured Monthly Transaction Data", 1)
    add_doc_bullets(
        document,
        [
            "Patient ID, patient name, visit date, doctor, PRO, billable item, modality, price, total discount, net amount.",
            "Total payment received, payment method, revenue booked in Sukhmani/Jivada, receipt status, notes, and raw source row JSON.",
            "Exact doctor incentive values when the monthly file contains Incentive to Doctors/RMPs or incentive workbook fields.",
        ],
    )

    add_doc_heading(document, "11. Calculation Logic", 1)
    add_doc_bullets(
        document,
        [
            "Allowed Discount is shown only when doctor, doctor group, and item mapping are all available.",
            "Incentive To Doctors uses exact incentive workbook value first, then incentive amount fallback, then max(allowed - actual, 0).",
            "Sum Of Both equals Total Discount Amount plus Incentive To Doctors.",
            "Variance equals Allowed minus Sum Of Both.",
            "Group A expects discount and incentive to be zero.",
            "Group B calculates allowed using Group D values and expects doctor incentive to be zero.",
            "Group C calculates allowed using Group F values and expects doctor incentive to be zero.",
            "Group rule violations force a red variance badge regardless of positive or negative value.",
        ],
    )

    add_doc_heading(document, "12. Calculation Trace Step-by-Step", 1)
    document.add_paragraph(
        "When the Run Engine button is clicked, the calculation follows this exact path from input rows to final remarks."
    )
    add_doc_table(document, ["Step", "Action", "Why It Happens", "Main Data Used"], CALCULATION_TRACE)

    add_doc_heading(document, "13. Why a Row Shows Each Remark", 1)
    add_doc_table(document, ["Remark", "Why It Shows", "User Action"], REMARK_EXPLAINER)

    add_doc_heading(document, "14. Latest Run Exception Summary", 1)
    add_doc_table(document, ["Remark", "Rows"], [[remark, str(rows)] for remark, rows in snapshot.latest_remarks])

    add_doc_heading(document, "15. Approval Workflow", 1)
    add_doc_bullets(
        document,
        [
            "Doctor info changes, PRO changes, doctor additions, incentive overrides, and disbursal approvals are stored as approval_requests.",
            "Pending requests do not directly change master/payment data.",
            "Approved doctor info requests update doctor_master.",
            "Approved PRO requests update doctor ownership mapping.",
            "Approved incentive override requests update payment amount.",
        ],
    )
    add_doc_table(document, ["Approval Type", "Status", "Rows"], [[typ, status, str(rows)] for typ, status, rows in snapshot.approvals])

    add_doc_heading(document, "16. Role-Based Working", 1)
    add_doc_table(document, ["Role", "Allowed Work", "Business Responsibility"], ROLE_GUIDE)

    add_doc_heading(document, "17. Button and Action Reference Appendix", 1)
    document.add_paragraph(
        "This appendix is only for detailed reference after the step workflow is understood. It explains what each main button/action triggers."
    )
    add_doc_table(document, ["Page", "Button or Action", "API or Trigger", "Database Area", "What It Does"], PAGE_ACTIONS)

    add_doc_heading(document, "18. Start-to-End Monthly Cycle", 1)
    add_doc_bullets(
        document,
        [
            "Start: Admin uploads reference master and requirement sheet in Setup Center.",
            "Next: User uploads monthly transaction file and incentive data in Monthly Intake.",
            "Calculation: Admin runs Calculation Review for the selected month.",
            "Review: Flags are checked for missing doctor, missing group, missing item, over-discount, and group policy mismatches.",
            "Approval: Required changes are sent through approval workflow.",
            "Payout: Valid results generate payout entries and cash-in-hand checks block disbursal when needed.",
            "Close: Reports are exported and the month is locked after completion.",
        ],
    )

    document.save(DOC_OUTPUT)


def image_for_pdf(path: Path, max_width: float, max_height: float) -> Image | None:
    if not path.exists():
        return None
    with PILImage.open(path) as img:
        width, height = img.size
    scale = min(max_width / width, max_height / height)
    return Image(str(path), width=width * scale, height=height * scale)


def build_pdf(snapshot: DbSnapshot) -> None:
    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    page_size = landscape(A4)
    doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=page_size,
        leftMargin=0.35 * inch,
        rightMargin=0.35 * inch,
        topMargin=0.35 * inch,
        bottomMargin=0.35 * inch,
        title="RRCP Architecture and Working Guide",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=26, textColor=colors.HexColor("#192b54"))
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=colors.HexColor("#192b54"), spaceBefore=8, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=colors.HexColor("#192b54"), spaceBefore=6, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, spaceAfter=4)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=5)
    small = ParagraphStyle("Small", parent=body, fontSize=6.6, leading=8, spaceAfter=2)
    code_style = ParagraphStyle("Code", parent=body, fontName="Courier", fontSize=5.8, leading=7, textColor=colors.HexColor("#1f2937"))

    story = [
        Paragraph("RRCP Architecture and Working Guide", title_style),
        Paragraph("End-to-end explanation with live screenshots, imports, data capture, calculation logic, approvals, outputs, and monthly close flow.", body),
        Spacer(1, 0.08 * inch),
    ]

    def add_heading(text: str, level: int = 1) -> None:
        story.append(Paragraph(text, h1 if level == 1 else h2))

    def add_bullets(items: list[str]) -> None:
        for item in items:
            story.append(Paragraph(item, bullet, bulletText="-"))

    def add_table(headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
        def cell(value: object, header: bool = False) -> Paragraph:
            style = h2 if header else small
            return Paragraph(escape(str(value)), style)

        table_data = [[cell(header, True) for header in headers]] + [
            [cell(value) for value in row] for row in rows
        ]
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3fb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#192b54")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.08 * inch))

    def add_preformatted(text: str) -> None:
        story.append(Preformatted(text, code_style, maxLineLength=155))
        story.append(Spacer(1, 0.08 * inch))

    def add_workflow_steps() -> None:
        for item in MONTHLY_WORKFLOW_STEPS:
            story.append(PageBreak())
            add_heading(f"{item['step']}: {item['title']}")
            add_table(
                ["Part", "Explanation"],
                [
                    ["Screen used", item["screen"]],
                    ["Where this step starts", item["starts_from"]],
                    ["Input used", item["input"]],
                    ["Data captured", item["data_captured"]],
                    ["Stored in", item["stored_in"]],
                    ["Why this step is needed", item["why"]],
                    ["Output of this step", item["output"]],
                ],
                [1.7 * inch, 8.4 * inch],
            )
            img = image_for_pdf(SCREENSHOT_DIR / item["screenshot"], 10.1 * inch, 3.15 * inch)
            if img:
                story.append(img)

    add_heading("1. Purpose and Stack")
    add_bullets(
        [
            "Purpose: manage monthly referral revenue, doctor incentives, approvals, payouts, reports, and month close.",
            "Frontend: HTML, CSS, vanilla JavaScript.",
            "Backend: PHP APIs running on XAMPP.",
            "Database: SQLite at data/app.db.",
            "Imports: XLSX/XLS reference, transaction, and incentive files.",
        ]
    )

    add_heading("2. Live Data Snapshot")
    add_table(["Table", "Rows"], [[name, str(count)] for name, count in snapshot.counts], [2.4 * inch, 1.0 * inch])

    latest_run = snapshot.latest_run
    summary = latest_run.get("summary", {})
    add_heading("3. Current Run")
    add_table(
        ["Metric", "Value"],
        [
            ["Run ID", str(latest_run.get("id", "-"))],
            ["Period", f"{latest_run.get('period_year', '-')}-{str(latest_run.get('period_month', '-')).zfill(2)}"],
            ["Total records", str(summary.get("totalRecords", latest_run.get("total_records", "-")))],
            ["Total flags", str(summary.get("totalFlags", latest_run.get("total_flags", "-")))],
            ["Total payable incentive", f"INR {summary.get('totalPayable', '-')}"],
            ["Run timestamp", str(latest_run.get("run_at", "-"))],
        ],
        [2.4 * inch, 3.2 * inch],
    )

    add_heading("4. Imports and Data Flow")
    add_table(["Upload Type", "File", "Rows", "Uploaded At"], [[typ, name, str(rows), at] for typ, name, rows, at in snapshot.latest_uploads])
    add_bullets(
        [
            "Setup Center -> reference tables: service_prices, discount_rules, doctor_master, software_requirements.",
            "Monthly Intake -> transactions and transaction_incentives.",
            "Calculation Review -> engine_runs and engine_results.",
            "Payout Center -> payments, approval_requests, pro_wallets, locked_periods.",
            "Reports -> CSV/PDF outputs using stored result tables.",
        ]
    )

    add_heading("5. Step-by-Step Working From Start to Final Output")
    story.append(Paragraph("The next pages explain the tool in actual operating order: prepare masters, verify doctors, upload monthly rows, calculate, resolve, approve, payout, report, and close.", body))
    add_workflow_steps()

    story.append(PageBreak())
    add_heading("6. Tree View of the Same Flow")
    story.append(Paragraph("This tree is included after the workflow as a compact technical view of the same journey.", body))
    add_preformatted(SYSTEM_TREE)

    story.append(PageBreak())
    add_heading("7. Field Data Lineage")
    story.append(Paragraph("This answers where each visible field is fetched from, where it is stored, and why the tool needs it.", body))
    add_table(
        ["Data Shown", "Fetched or Captured From", "Stored In", "Why It Is Needed"],
        DATA_LINEAGE,
        [1.55 * inch, 2.35 * inch, 2.35 * inch, 3.8 * inch],
    )

    story.append(PageBreak())
    add_heading("8. Calculation Rules")
    add_bullets(
        [
            "Allowed is computed only when doctor, group, and item are all available.",
            "Incentive uses exact imported value first, then fallback logic.",
            "Variance = Allowed - (Total Discount Amount + Incentive To Doctors).",
            "Group A expects both discount and incentive to be zero.",
            "Group B uses Group D allowed discount and zero doctor incentive.",
            "Group C uses Group F allowed discount and zero doctor incentive.",
            "Group rule mismatch forces red variance.",
        ]
    )
    add_table(
        ["Step", "Action", "Why It Happens", "Main Data Used"],
        CALCULATION_TRACE,
        [0.45 * inch, 1.4 * inch, 5.1 * inch, 3.05 * inch],
    )

    add_heading("9. Why Rows Show Each Remark")
    add_table(
        ["Remark", "Why It Shows", "User Action"],
        REMARK_EXPLAINER,
        [2.3 * inch, 4.3 * inch, 3.45 * inch],
    )

    add_heading("10. Current Run Exception Summary")
    add_table(["Remark", "Rows"], [[remark, str(rows)] for remark, rows in snapshot.latest_remarks])

    add_heading("11. Approval State")
    add_table(["Approval Type", "Status", "Rows"], [[typ, status, str(rows)] for typ, status, rows in snapshot.approvals])

    add_heading("12. Role-Based Working")
    add_table(["Role", "Allowed Work", "Business Responsibility"], ROLE_GUIDE, [1.3 * inch, 4.6 * inch, 4.1 * inch])

    story.append(PageBreak())
    add_heading("13. Button and Action Reference Appendix")
    story.append(Paragraph("This appendix is kept as supporting detail. The primary explanation is the step workflow above.", body))
    add_table(
        ["Page", "Button or Action", "API or Trigger", "Database Area", "What It Does"],
        PAGE_ACTIONS,
        [1.0 * inch, 1.25 * inch, 1.55 * inch, 1.65 * inch, 4.65 * inch],
    )

    story.append(PageBreak())
    add_heading("8. Start-to-End Cycle")
    add_bullets(
        [
            "Upload reference master and requirement sheet.",
            "Verify doctor mapping, group, PRO, cycle, reporting, confirmation, and verified status.",
            "Upload monthly transaction and incentive data.",
            "Run Calculation Review.",
            "Resolve flags and approvals.",
            "Generate payouts and confirm cash-in-hand status.",
            "Export reports and lock the period.",
        ]
    )

    doc.build(story)


def main() -> None:
    snapshot = fetch_snapshot()
    build_docx(snapshot)
    build_pdf(snapshot)
    print(DOC_OUTPUT)
    print(PDF_OUTPUT)


if __name__ == "__main__":
    main()
